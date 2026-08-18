from __future__ import annotations

import math
import re
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Protocol

from .models import CalculationResult, MicropileInput, PileType
from .symbols import SYMBOLS, EngineeringSymbol
from .schematic_image import render_pile_soil_schematic


class ReportGenerator(Protocol):
    def generate(self, result: CalculationResult, output_path: Path) -> Path: ...


REPORT_SYMBOLS: dict[str, EngineeringSymbol] = SYMBOLS


class DocxReportGenerator:
    """Generate an A4 calculation book from the normalized calculation result."""

    TABLE_WIDTH_DXA = 9072
    BLUE = "1F4D78"
    LIGHT_BLUE = "E8EEF5"
    LIGHT_GRAY = "F2F4F7"
    GREEN = "147D3F"
    RED = "B42318"

    def generate(self, result: CalculationResult, output_path: Path) -> Path:
        if result.normalized_input is None:
            raise ValueError("计算结果缺少规范化输入，无法生成计算书")
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Mm, Pt, RGBColor

        data = result.normalized_input
        doc = Document()
        section = doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(22)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)

        styles = doc.styles
        math_pr = doc.settings._element.find(qn("m:mathPr"))
        if math_pr is not None:
            math_font = math_pr.find(qn("m:mathFont"))
            if math_font is not None:
                math_font.set(qn("m:val"), "Times New Roman")
        normal = styles["Normal"]
        normal.font.name = "Times New Roman"
        normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        normal._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        normal.font.size = Pt(12)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25
        for style_name, size, before, after in (
            ("Heading 1", 14, 18, 10),
            ("Heading 2", 14, 14, 7),
            ("Heading 3", 14, 10, 5),
        ):
            style = styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            for theme_attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
                style._element.rPr.rFonts.attrib.pop(qn(f"w:{theme_attribute}"), None)
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(self.BLUE)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._add_text_run(header, "光伏支架微型桩规范验算计算书", 9, color="5F6B7A")
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_text_run(footer, "第 ", 9, color="5F6B7A")
        self._add_field(footer, "PAGE")
        self._add_text_run(footer, " 页", 9, color="5F6B7A")

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(28)
        title.paragraph_format.space_after = Pt(8)
        self._add_text_run(title, "光伏支架微型桩规范验算计算书", 22, bold=True)
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(22)
        self._add_text_run(subtitle, data.project_name or "未命名项目", 14, bold=True, color=self.BLUE)
        metadata = [
            ("项目名称", data.project_name or "未命名项目"),
            ("桩型", data.pile_type.value),
            ("计算日期", datetime.now().astimezone().strftime("%Y-%m-%d")),
            ("计算依据", "NB/T 10115-2018、GB 51101-2016、JGJ 94-2008"),
        ]
        self._add_table(doc, ["项目", "内容"], metadata, [1900, 7172])

        doc.add_heading("1 计算条件", level=1)
        self._add_input_sections(doc, data)
        doc.add_heading("1.4 桩土示意图", level=2)
        self._add_schematic(doc, data)
        doc.add_heading("2 验算结果汇总", level=1)
        summary_rows = []
        for check in result.checks.values():
            capacity = "无需验算" if math.isinf(check.capacity_kn) else f"{check.capacity_kn:.3f} kN"
            utilization = "—" if math.isinf(check.utilization) else f"{check.utilization:.3f}"
            summary_rows.append((
                check.name,
                f"{check.demand_kn:.3f} kN",
                capacity,
                utilization,
                "满足" if check.passed else "不满足",
                check.clause,
            ))
        self._add_table(
            doc,
            ["验算项目", "作用值", "承载力", "利用率", "结论", "规范依据"],
            summary_rows,
            [1500, 1100, 1200, 850, 850, 3572],
            status_column=4,
        )

        doc.add_heading("3 主要计算过程", level=1)
        self._add_compression_calculation(doc, result)
        self._add_uplift_calculation(doc, result)
        self._add_horizontal_calculation(doc, result)
        self._add_stability_calculation(doc, result)

        doc.add_heading("4 结论", level=1)
        failed = [check.name for check in result.checks.values() if not check.passed]
        conclusion = (
            f"本项目微型桩的抗压、抗拔、水平承载力及整体稳定四项验算均满足要求。"
            if not failed
            else f"本项目存在未通过验算项目：{'、'.join(failed)}，应调整设计参数后重新验算。"
        )
        paragraph = doc.add_paragraph()
        self._add_text_run(paragraph, conclusion, 12, bold=True, color=self.GREEN if not failed else self.RED)

        self._enforce_heading_fonts(doc)

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return output_path

    def _enforce_heading_fonts(self, doc) -> None:
        sizes = {"Heading 1": 14, "Heading 2": 14, "Heading 3": 14}
        for paragraph in doc.paragraphs:
            size = sizes.get(paragraph.style.name)
            if size is None:
                continue
            for run in paragraph.runs:
                self._format_run(run, size, bold=True, color=self.BLUE)

    def _add_input_sections(self, doc, data: MicropileInput) -> None:
        common = data.common
        load_rows = [
            ("桩顶压力 N_Mk", f"{data.loads.compression_kn:g}", "kN"),
            ("桩顶拔力 T_k", f"{data.loads.uplift_kn:g}", "kN"),
            ("桩顶水平力 H_Mik", f"{data.loads.horizontal_kn:g}", "kN"),
            ("抗拔验算考虑桩自重 G_p", "是" if data.loads.consider_pile_self_weight else "否", "—"),
        ]
        doc.add_heading("1.1 荷载与通用参数", level=2)
        common_rows = [
            ("桩身外径 d", f"{common.diameter_m * 1000:g}", "mm"),
            ("地下计算埋深 h_t", f"{common.embedment_m:g}", "m"),
            ("桩高出地面高度 h_0", f"{common.above_ground_height_m:g}", "m"),
            ("桩顶约束", common.top_constraint.value, "—"),
            ("允许地面处水平位移 x_0a", f"{common.allowable_displacement_mm:g}", "mm"),
            ("水平抗力比例系数 m", f"{common.horizontal_m_mn_m4:g}", "MN/m⁴"),
            ("JGJ 94水平计算土类", common.horizontal_soil_class, "—"),
            ("附录C桩端条件", common.pile_tip_condition.value, "—"),
            ("整体稳定土类", common.stability_soil_type.value, "—"),
            *load_rows,
        ]
        if common.rock_strength_kpa is not None:
            insertion_index = len(common_rows) - len(load_rows) - 1
            common_rows.insert(insertion_index, ("岩石饱和单轴抗压强度标准值 f_rk", f"{common.rock_strength_kpa:g}", "kPa"))
        self._add_table(doc, ["参数", "输入值", "单位"], common_rows, [4300, 3272, 1500])

        doc.add_heading("1.2 桩截面参数", level=2)
        if data.pile_type is PileType.GROUTED:
            assert data.grouted is not None
            section_rows = [
                ("混凝土弹性模量 E_c", f"{data.grouted.concrete_modulus_mpa:g}", "MPa"),
                ("钢筋弹性模量 E_s", f"{data.grouted.steel_modulus_mpa:g}", "MPa"),
                ("现有配筋率 ρ_g", f"{data.grouted.reinforcement_ratio * 100:g}", "%"),
                ("保护层厚度 c", f"{data.grouted.cover_m * 1000:g}", "mm"),
            ]
        else:
            assert data.helical is not None
            section_rows = [
                ("钢管壁厚 t", f"{data.helical.wall_thickness_m * 1000:g}", "mm"),
                ("钢材弹性模量 E", f"{data.helical.steel_modulus_mpa:g}", "MPa"),
                ("叶片直径 D", f"{data.helical.blade_diameter_m * 1000:g}", "mm"),
                ("叶片埋深", "、".join(f"{depth:g}" for depth in data.helical.blade_depths_m), "m"),
            ]
        self._add_table(doc, ["参数", "输入值", "单位"], section_rows, [4300, 3272, 1500])

        doc.add_heading("1.3 地质参数", level=2)
        soil_rows = [
            (layer.name, f"{layer.thickness_m:g}", f"{layer.unit_weight_kn_m3:g}", f"{layer.beta_deg:g}",
             f"{layer.qsik_kpa:g}", f"{layer.qpk_kpa:g}", f"{layer.uplift_factor:g}")
            for layer in data.soils
        ]
        self._add_table(
            doc,
            ["土层", "厚度/m", "γ/(kN/m³)", "β/(°)", "q_sik/kPa", "q_pk/kPa", "λ"],
            soil_rows,
            [1500, 900, 1300, 900, 1600, 1600, 1272],
        )

    def _add_compression_calculation(self, doc, result: CalculationResult) -> None:
        doc.add_heading("3.1 抗压验算", level=2)
        paragraph = doc.add_paragraph()
        if result.pile_type is PileType.GROUTED:
            self._add_text_run(paragraph, "微型灌注桩抗压承载力由桩侧极限阻力和桩端极限阻力组成，按下式计算：", 12)
            self._formula(doc, [REPORT_SYMBOLS["Q_UK"], " = πdΣ", EngineeringSymbol("q", "sik"), EngineeringSymbol("l", "i"), " + ", EngineeringSymbol("q", "pk"), "πd²/4"])
            self._formula(doc, [REPORT_SYMBOLS["Q_SK"], " = πdΣ", EngineeringSymbol("q", "sik"), EngineeringSymbol("l", "i")])
            self._process_value(doc, "桩侧极限阻力", REPORT_SYMBOLS["Q_SK"], self._number(result, "竖向总极限侧阻力 Qsk (kN)"), "kN")
            self._formula(doc, [REPORT_SYMBOLS["Q_PK"], " = ", EngineeringSymbol("q", "pk"), "πd²/4"])
            self._process_value(doc, "桩端极限阻力", REPORT_SYMBOLS["Q_PK"], self._number(result, "桩端极限阻力 Qpk (kN)"), "kN")
        else:
            self._add_text_run(paragraph, "钢螺旋桩按GB 51101-2016式（5.3.9）计算，桩周计算周长按表5.3.9根据叶片位置分段取值。", 12)
            self._formula(doc, [REPORT_SYMBOLS["Q_UK"], " = Σ", EngineeringSymbol("u", "ci"), EngineeringSymbol("q", "sik"), EngineeringSymbol("l", "i"), " + ", EngineeringSymbol("q", "pk"), EngineeringSymbol("A", "D")])
            self._formula(doc, [EngineeringSymbol("A", "D"), " = π(D²-d²)/4"])
            self._process_value(doc, "叶片净投影面积", EngineeringSymbol("A", "D"), self._number(result, "叶片净投影面积 AD (m²)"), "m²")
            self._process_value(doc, "有效侧阻力", EngineeringSymbol("Q", "s"), self._number(result, "螺旋桩有效侧阻力 (kN)"), "kN")
            self._process_value(doc, "叶片端阻力", EngineeringSymbol("Q", "p"), self._number(result, "叶片端阻力 (kN)"), "kN")
        self._process_value(doc, "抗压极限承载力", REPORT_SYMBOLS["Q_UK"], self._number(result, "抗压极限承载力 Quk (kN)"), "kN")
        self._formula(doc, [EngineeringSymbol("R", "c"), " = ", REPORT_SYMBOLS["Q_UK"], "/K，K=2"])
        self._process_value(doc, "抗压承载力", EngineeringSymbol("R", "c"), result.checks["compression"].capacity_kn, "kN")
        self._result_sentence(doc, result.checks["compression"])

    def _add_uplift_calculation(self, doc, result: CalculationResult) -> None:
        doc.add_heading("3.2 抗拔验算", level=2)
        paragraph = doc.add_paragraph()
        if result.pile_type is PileType.GROUTED:
            self._add_text_run(paragraph, "微型灌注桩按GB 51101-2016式（5.3.8）计算各土层抗拔侧阻之和：", 12)
            self._formula(doc, [REPORT_SYMBOLS["T_UK"], " = πdΣλ", EngineeringSymbol("q", "sik"), EngineeringSymbol("l", "i")])
        else:
            self._add_text_run(paragraph, "钢螺旋桩按GB 51101-2016式（5.3.10）计算，抗拔桩周计算周长按表5.3.10根据叶片位置及间距分段取值：", 12)
            self._formula(doc, [REPORT_SYMBOLS["T_UK"], " = Σλ", EngineeringSymbol("u", "ti"), EngineeringSymbol("q", "sik"), EngineeringSymbol("l", "i")])
        paragraph = doc.add_paragraph()
        self._add_text_run(paragraph, "各土层抗拔系数按第5.3.8条取值：岩石λ=0.8，砂土λ=0.5，黏性土或粉土λ=0.7。", 12)
        if result.normalized_input.loads.consider_pile_self_weight:
            if result.pile_type is PileType.GROUTED:
                self._add_text_run(
                    doc.add_paragraph(),
                    "勾选考虑桩自重。微型灌注桩按钢筋混凝土圆形实心截面计算，桩身总长度取地上长度与地下埋深之和，材料重度取25 kN/m³：",
                    12,
                )
                self._formula(doc, [EngineeringSymbol("V", "p"), " = πd²(", EngineeringSymbol("h", "0"), "+", EngineeringSymbol("h", "t"), ")/4"])
                self._formula(doc, [EngineeringSymbol("G", "p"), " = 25", EngineeringSymbol("V", "p")])
            else:
                self._add_text_run(
                    doc.add_paragraph(),
                    "勾选考虑桩自重。钢螺旋桩按钢管空心圆截面计算，桩身总长度取地上长度与地下埋深之和，钢材重度取78 kN/m³，忽略螺旋叶片重量：",
                    12,
                )
                self._formula(doc, [EngineeringSymbol("V", "p"), " = π[d²-(d-2t)²](", EngineeringSymbol("h", "0"), "+", EngineeringSymbol("h", "t"), ")/4"])
                self._formula(doc, [EngineeringSymbol("G", "p"), " = 78", EngineeringSymbol("V", "p")])
            self._process_value(doc, "桩自重计算体积", EngineeringSymbol("V", "p"), self._number(result, "桩自重计算体积 Vp (m³)"), "m³")
            self._process_value(doc, "单桩自重", EngineeringSymbol("G", "p"), self._number(result, "抗拔验算采用单桩自重 Gp (kN)"), "kN")
        else:
            self._formula(doc, ["本程序保守取单桩自重 ", EngineeringSymbol("G", "p"), " = 0。"])
        self._process_value(doc, "抗拔极限承载力", REPORT_SYMBOLS["T_UK"], self._number(result, "抗拔极限承载力 Tuk (kN)"), "kN")
        self._formula(doc, [EngineeringSymbol("R", "t"), " = ", REPORT_SYMBOLS["T_UK"], "/K，K=2"])
        self._process_value(doc, "抗拔承载力", EngineeringSymbol("R", "t"), result.checks["uplift"].capacity_kn, "kN")
        self._formula(doc, [EngineeringSymbol("R", "t"), " ≥ max(", REPORT_SYMBOLS["T_K"], " - ", EngineeringSymbol("G", "p"), ", 0)"])
        paragraph = doc.add_paragraph()
        self._add_text_run(
            paragraph,
            (
                f"桩顶拔力={result.normalized_input.loads.uplift_kn:.3f} kN，"
                f"采用桩自重={self._number(result, '抗拔验算采用单桩自重 Gp (kN)'):.3f} kN，"
                f"抗拔验算净作用={result.checks['uplift'].demand_kn:.3f} kN，"
                f"抗拔承载力={result.checks['uplift'].capacity_kn:.3f} kN，"
                f"利用率={result.checks['uplift'].utilization:.3f}，"
                f"验算结论：{'满足' if result.checks['uplift'].passed else '不满足'}。"
            ),
            12,
        )

    def _add_schematic(self, doc, data: MicropileInput) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm

        with tempfile.TemporaryDirectory(prefix="micropile_report_") as directory:
            image_path = render_pile_soil_schematic(data, Path(directory) / "pile_soil_schematic.png")
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = self._pt(4)
            shape = paragraph.add_run().add_picture(str(image_path), width=Mm(150))
            shape._inline.docPr.set("descr", "根据输入参数按竖向比例绘制的桩土示意图")
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = self._pt(8)
        self._add_text_run(caption, "图1  桩土示意图（竖向尺寸按输入比例绘制，桩身横向宽度仅作示意）", 10.5, color="5F6B7A")

    def _add_horizontal_calculation(self, doc, result: CalculationResult) -> None:
        doc.add_heading("3.3 水平承载力验算", level=2)
        if math.isinf(result.checks["horizontal"].capacity_kn):
            self._result_sentence(doc, result.checks["horizontal"])
            return
        if result.pile_type is PileType.GROUTED:
            self._process_value(doc, "扣除保护层后的截面直径", EngineeringSymbol("d", "0"), self._number(result, "扣除保护层直径 d0 (m)"), "m")
            self._process_value(doc, "钢筋与混凝土弹性模量比", EngineeringSymbol("α", "E"), self._number(result, "钢筋与混凝土弹模比 αE"))
            self._process_value(doc, "换算截面模量", EngineeringSymbol("W", "0"), self._number(result, "换算截面模量 W0 (m³)"), "m³")
            self._process_value(doc, "换算截面惯性矩", EngineeringSymbol("I", "0"), self._number(result, "换算惯性矩 I0 (m⁴)"), "m⁴")
        else:
            self._process_value(doc, "钢管内径", EngineeringSymbol("d", "i"), self._number(result, "钢管内径 (m)"), "m")
            self._process_value(doc, "钢管截面惯性矩", EngineeringSymbol("I"), self._number(result, "钢管截面惯性矩 I (m⁴)"), "m⁴")
        self._process_value(doc, "桩身抗弯刚度", EngineeringSymbol("EI"), self._number(result, "桩身抗弯刚度 EI (kN·m²)"), "kN·m²")
        self._process_value(doc, "JGJ 94基础计算宽度", EngineeringSymbol("b", "0,base"), self._number(result, "JGJ基础计算宽度 b0,base (m)"), "m")
        self._process_value(doc, "折减后计算宽度", REPORT_SYMBOLS["B_0"], self._number(result, "折减后计算宽度 b0 (m)"), "m")
        self._formula(doc, ["α = (m", REPORT_SYMBOLS["B_0"], "/EI)^(1/5)"])
        self._process_value(doc, "水平变形系数", EngineeringSymbol("α"), self._number(result, "水平变形系数 α (1/m)"), "1/m")
        self._process_value(doc, "换算埋深", EngineeringSymbol("αh"), self._number(result, "换算埋深 αh"))
        paragraph = doc.add_paragraph()
        self._add_text_run(
            paragraph,
            "本项按JGJ 94-2008附录C表C.0.3-1有限长度桩m法计算地面处水平位移x₀；"
            "桩高出地面段通过地面处内力参与计算，其桩顶附加位移不计入基础位移限值。",
            12,
        )
        self._add_text_run(
            paragraph,
            f"桩端条件采用“{result.intermediates['附录C桩端条件']}”。",
            12,
        )
        c0 = self._number(result, "桩底竖向抗力系数 C0 (kN/m³)")
        if math.isfinite(c0):
            self._process_value(doc, "桩底竖向抗力系数", EngineeringSymbol("C", "0"), c0, "kN/m³")
        else:
            paragraph = doc.add_paragraph()
            self._add_text_run(paragraph, "桩端嵌固于基岩，按表C.0.3-1注2采用嵌固边界条件。", 12)
        self._process_value(doc, "桩端约束系数", EngineeringSymbol("K", "h"), self._number(result, "附录C桩端约束系数 Kh"))
        self._process_value(doc, "标准组合水平力传至地面处的弯矩", EngineeringSymbol("M", "0k"), self._number(result, "标准组合水平力下地面处弯矩 M0k (kN·m)"), "kN·m")
        self._process_value(doc, "单位水平力地面处位移", EngineeringSymbol("δ", "x0"), self._number(result, "单位水平力地面处位移 δx0 (m/kN)"), "m/kN")
        self._formula(doc, [EngineeringSymbol("x", "0k"), " = 1000", EngineeringSymbol("δ", "x0"), "·", REPORT_SYMBOLS["H_MIK"]])
        self._process_value(doc, "标准组合水平力下地面处位移", EngineeringSymbol("x", "0k"), self._number(result, "标准组合水平力下地面处位移 x0k (mm)"), "mm")
        self._formula(doc, [REPORT_SYMBOLS["R_HA"], " = 0.75(", REPORT_SYMBOLS["X_0A"], "/1000)/", EngineeringSymbol("δ", "x0")])
        self._process_value(doc, "水平承载力特征值", REPORT_SYMBOLS["R_HA"], self._number(result, "水平承载力特征值 Rha (kN)"), "kN")
        self._result_sentence(doc, result.checks["horizontal"])

    def _add_stability_calculation(self, doc, result: CalculationResult) -> None:
        doc.add_heading("3.4 整体稳定（抗倾覆）验算", level=2)
        if math.isinf(result.checks["stability"].capacity_kn):
            self._result_sentence(doc, result.checks["stability"])
            return
        self._process_value(doc, "桩埋深范围内加权平均重度", EngineeringSymbol("γ", "s"), self._number(result, "加权平均重度 γs (kN/m³)"), "kN/m³")
        self._process_value(doc, "桩埋深范围内加权等代内摩擦角", EngineeringSymbol("β"), self._number(result, "加权等代内摩擦角 β (°)"), "°")
        self._process_value(doc, "水平力高度比", EngineeringSymbol("η"), self._number(result, "水平力高度比 η"))
        self._process_value(doc, "土的侧压力系数（表8.3.15-3）", EngineeringSymbol("ξ"), self._number(result, "土的侧压力系数 ξ（表8.3.15-3）"))
        self._process_value(doc, "式8.3.15-9求得压力扩散角参数", EngineeringSymbol("θ"), self._number(result, "压力扩散角参数 θ"))
        self._process_value(doc, "接触面间摩阻系数", EngineeringSymbol("μ", "cm"), self._number(result, "接触面间摩阻系数 μcm"))
        self._process_value(doc, "土压力参数", EngineeringSymbol("m"), self._number(result, "土压力参数 m (kN/m³)"), "kN/m³")
        self._process_value(doc, "空间增大系数", EngineeringSymbol("K", "0"), self._number(result, "空间增大系数 K0"))
        self._process_value(doc, "整体稳定计算宽度", REPORT_SYMBOLS["B_0"], self._number(result, "整体稳定计算宽度 b0 (m)"), "m")
        self._formula(doc, [REPORT_SYMBOLS["R_H"], " = m", REPORT_SYMBOLS["B_0"], EngineeringSymbol("h", "t"), "²/(η", EngineeringSymbol("μ", "cm"), ")"])
        self._process_value(doc, "整体稳定水平抗力", REPORT_SYMBOLS["R_H"], self._number(result, "整体稳定水平抗力 RH (kN)"), "kN")
        self._process_value(doc, "整体稳定验算系数", REPORT_SYMBOLS["K_MW"], 1.1)
        self._formula(
            doc,
            [
                REPORT_SYMBOLS["R_H"],
                " ≥ ",
                REPORT_SYMBOLS["K_MW"],
                "×",
                REPORT_SYMBOLS["H_MIK"],
                " = 1.1",
                REPORT_SYMBOLS["H_MIK"],
            ],
        )
        self._result_sentence(doc, result.checks["stability"])

    @staticmethod
    def _number(result: CalculationResult, key: str) -> float:
        value = result.intermediates.get(key)
        if not isinstance(value, (int, float)):
            raise ValueError(f"计算书缺少中间计算值：{key}")
        return float(value)

    def _process_value(self, doc, description: str, symbol: EngineeringSymbol, value: float, unit: str = "") -> None:
        suffix = f" {unit}" if unit else ""
        self._formula(doc, [f"{description}：", symbol, f" = {value:.6g}{suffix}"])

    def _result_sentence(self, doc, check) -> None:
        capacity = "无需验算" if math.isinf(check.capacity_kn) else f"{check.capacity_kn:.3f} kN"
        text = f"作用值={check.demand_kn:.3f} kN，承载力={capacity}，"
        if not math.isinf(check.utilization):
            text += f"利用率={check.utilization:.3f}，"
        text += f"验算结论：{'满足' if check.passed else '不满足'}。"
        paragraph = doc.add_paragraph()
        self._add_text_run(paragraph, text, 12)

    def _formula(self, doc, parts: list[str | EngineeringSymbol]) -> None:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = self._mm(8)
        paragraph.paragraph_format.space_after = self._pt(4)
        for part in parts:
            if isinstance(part, EngineeringSymbol):
                for value, subscript in part.runs:
                    run = paragraph.add_run(value)
                    self._format_run(run, 12)
                    run.font.subscript = subscript
            else:
                self._add_text_run(paragraph, part, 12)

    def _add_table(self, doc, headers, rows, widths, status_column: int | None = None):
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        table = doc.add_table(rows=1, cols=len(headers))
        table.autofit = False
        table.style = "Table Grid"
        self._set_table_geometry(table, widths)
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            self._set_cell_text(cell, str(header), 10.5, bold=True)
            self._shade_cell(cell, self.LIGHT_BLUE)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)
        for row_values in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row_values):
                color = None
                bold = False
                if status_column == index:
                    color = self.GREEN if str(value) == "满足" else self.RED
                    bold = True
                self._set_cell_text(cells[index], str(value), 10.5, bold=bold, color=color)
                cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cells[index].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index in (0, len(headers) - 1) else WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_after = self._pt(0)
        self._set_table_geometry(table, widths)
        doc.add_paragraph().paragraph_format.space_after = self._pt(2)
        return table

    def _set_cell_text(self, cell, text: str, size: float, bold: bool = False, color: str | None = None) -> None:
        paragraph = cell.paragraphs[0]
        paragraph.clear()
        pattern = re.compile(r"([A-Za-zα-ωΑ-Ω]+)_([A-Za-z0-9]+)")
        position = 0
        for match in pattern.finditer(text):
            if match.start() > position:
                self._add_text_run(paragraph, text[position:match.start()], size, bold=bold, color=color)
            base = paragraph.add_run(match.group(1))
            self._format_run(base, size, bold=bold, color=color)
            subscript = paragraph.add_run(match.group(2))
            self._format_run(subscript, size, bold=bold, color=color)
            subscript.font.subscript = True
            position = match.end()
        if position < len(text):
            self._add_text_run(paragraph, text[position:], size, bold=bold, color=color)

    def _set_table_geometry(self, table, widths) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        if sum(widths) != self.TABLE_WIDTH_DXA:
            raise ValueError("计算书表格列宽总和必须等于9072 DXA")
        table_pr = table._tbl.tblPr
        table_width = table_pr.first_child_found_in("w:tblW")
        table_width.set(qn("w:type"), "dxa")
        table_width.set(qn("w:w"), str(self.TABLE_WIDTH_DXA))
        table_indent = OxmlElement("w:tblInd")
        table_indent.set(qn("w:type"), "dxa")
        table_indent.set(qn("w:w"), "120")
        table_pr.append(table_indent)
        cell_margins = OxmlElement("w:tblCellMar")
        for edge, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")
            cell_margins.append(node)
        table_pr.append(cell_margins)
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            column = OxmlElement("w:gridCol")
            column.set(qn("w:w"), str(width))
            grid.append(column)
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                tc_width.set(qn("w:type"), "dxa")
                tc_width.set(qn("w:w"), str(width))

    @staticmethod
    def _shade_cell(cell, fill: str) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(shading)

    def _add_text_run(self, paragraph, text, size, bold=False, color=None):
        run = paragraph.add_run(text)
        self._format_run(run, size, bold=bold, color=color)
        return run

    @staticmethod
    def _format_run(run, size, bold=False, color=None) -> None:
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        run.font.name = "Times New Roman"
        r_fonts = run._element.get_or_add_rPr().rFonts
        r_fonts.set(qn("w:ascii"), "Times New Roman")
        r_fonts.set(qn("w:hAnsi"), "Times New Roman")
        r_fonts.set(qn("w:cs"), "Times New Roman")
        r_fonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    @staticmethod
    def _add_field(paragraph, instruction: str) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        run = paragraph.add_run()
        DocxReportGenerator._format_run(run, 9, color="5F6B7A")
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        text = OxmlElement("w:instrText")
        text.set(qn("xml:space"), "preserve")
        text.text = instruction
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        display = OxmlElement("w:t")
        display.text = "1"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend((begin, text, separate, display, end))

    @staticmethod
    def _pt(value):
        from docx.shared import Pt
        return Pt(value)

    @staticmethod
    def _mm(value):
        from docx.shared import Mm
        return Mm(value)
