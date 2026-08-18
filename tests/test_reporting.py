from __future__ import annotations

import sys
import tempfile
import unittest
import zipfile
from dataclasses import replace
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from micropile_app.calculations import calculate  # noqa: E402
from micropile_app.models import PileTipCondition  # noqa: E402
from micropile_app.reporting import DocxReportGenerator  # noqa: E402
from test_calculations import sample_grouted, sample_helical  # noqa: E402


class ReportingTests(unittest.TestCase):
    def _generate_and_read_xml(self, result):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "计算书.docx"
            DocxReportGenerator().generate(result, path)
            self.assertGreater(path.stat().st_size, 10_000)
            report = Document(path)
            heading_fonts_ok = all(
                all(
                    run._element.rPr is not None
                    and run._element.rPr.rFonts is not None
                    and run._element.rPr.rFonts.get(qn("w:ascii")) == "Times New Roman"
                    and run._element.rPr.rFonts.get(qn("w:eastAsia")) == "宋体"
                    and run.font.size is not None
                    and abs(run.font.size.pt - 14) < 1e-9
                    and run.bold is True
                    for run in paragraph.runs
                )
                for paragraph in report.paragraphs
                if paragraph.style.name.startswith("Heading")
            )
            normal_size_ok = abs(report.styles["Normal"].font.size.pt - 12) < 1e-9
            table_sizes_ok = all(
                run.font.size is not None and abs(run.font.size.pt - 10.5) < 1e-9
                for table in report.tables
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
                for run in paragraph.runs
                if run.text
            )
            with zipfile.ZipFile(path) as archive:
                xml = archive.read("word/document.xml").decode("utf-8")
                styles = archive.read("word/styles.xml").decode("utf-8")
                all_xml = "".join(
                    archive.read(name).decode("utf-8")
                    for name in archive.namelist()
                    if name.endswith(".xml")
                )
                media = [name for name in archive.namelist() if name.startswith("word/media/")]
        return xml, styles, all_xml, media, heading_fonts_ok, normal_size_ok, table_sizes_ok

    def test_grouted_report_contains_inputs_results_and_true_subscripts(self) -> None:
        result = calculate(replace(sample_grouted(), project_name="灌注桩测试项目"))
        xml, styles, all_xml, media, heading_fonts_ok, normal_size_ok, table_sizes_ok = self._generate_and_read_xml(result)
        self.assertIn("灌注桩测试项目", xml)
        self.assertIn("验算结果汇总", xml)
        self.assertIn("抗压验算", xml)
        self.assertIn("3.1 抗压验算", xml)
        self.assertIn("3.2 抗拔验算", xml)
        self.assertIn("3.3 水平承载力验算", xml)
        self.assertIn("3.4 整体稳定（抗倾覆）验算", xml)
        self.assertNotIn("3.1 竖向承载力验算", xml)
        self.assertIn("各土层抗拔系数按第5.3.8条取值", xml)
        self.assertIn("本程序保守取单桩自重", xml)
        self.assertIn("G", xml)
        self.assertIn("p", xml)
        self.assertIn('w:vertAlign w:val="subscript"', xml)
        self.assertNotIn("小直径折减", xml)
        self.assertEqual(len(media), 1)
        self.assertIn("桩土示意图", xml)
        self.assertNotIn("主要中间量", xml)
        self.assertNotIn("抗压控制土层", xml)
        self.assertNotIn("抗拔控制土层", xml)
        self.assertIn("附录C桩端条件", xml)
        self.assertIn("换算截面惯性矩", xml)
        self.assertIn("整体稳定水平抗力", xml)
        self.assertIn('w:ascii="Times New Roman"', xml + styles)
        self.assertIn('w:eastAsia="宋体"', xml + styles)
        self.assertNotIn("Microsoft YaHei", all_xml)
        self.assertNotIn("Cambria Math", all_xml)
        self.assertTrue(heading_fonts_ok)
        self.assertTrue(normal_size_ok)
        self.assertTrue(table_sizes_ok)

    def test_report_explains_selected_grouted_self_weight(self) -> None:
        base = sample_grouted()
        data = replace(base, loads=replace(base.loads, consider_pile_self_weight=True))
        result = calculate(data)
        xml, _styles, _all_xml, _media, *_checks = self._generate_and_read_xml(result)
        self.assertIn("勾选考虑桩自重", xml)
        self.assertIn("材料重度取25 kN/m³", xml)
        self.assertIn("抗拔验算净作用", xml)
        self.assertNotIn("本程序保守取单桩自重", xml)

    def test_report_explains_selected_helical_self_weight(self) -> None:
        base = sample_helical()
        data = replace(base, loads=replace(base.loads, consider_pile_self_weight=True))
        result = calculate(data)
        xml, _styles, _all_xml, _media, *_checks = self._generate_and_read_xml(result)
        self.assertIn("钢材重度取78 kN/m³", xml)
        self.assertIn("忽略螺旋叶片重量", xml)

    def test_helical_report_contains_blade_parameters(self) -> None:
        result = calculate(replace(sample_helical(), project_name="螺旋桩测试项目"))
        xml, _styles, _all_xml, media, _heading_fonts_ok, _normal_size_ok, _table_sizes_ok = self._generate_and_read_xml(result)
        self.assertIn("螺旋桩测试项目", xml)
        self.assertIn("叶片直径", xml)
        self.assertIn("叶片净投影面积", xml)
        self.assertIn("桩周计算周长按表5.3.9", xml)
        self.assertIn("抗拔桩周计算周长按表5.3.10", xml)
        self.assertEqual(len(media), 1)

    def test_appendix_c_report_contains_ground_displacement_process(self) -> None:
        base = sample_grouted()
        soils = (replace(base.soils[1], thickness_m=1.0),)
        common = replace(
            base.common,
            embedment_m=1.0,
            above_ground_height_m=0.5,
            horizontal_m_mn_m4=2.5,
            pile_tip_condition=PileTipCondition.ROCK_SURFACE,
            rock_strength_kpa=10000,
        )
        result = calculate(replace(base, soils=soils, common=common))
        xml, _styles, _all_xml, _media, *_checks = self._generate_and_read_xml(result)
        self.assertIn("附录C桩端条件", xml)
        self.assertIn("附录C表C.0.3-1有限长度桩m法", xml)
        self.assertIn("桩端条件采用", xml)
        self.assertIn("桩底竖向抗力系数", xml)
        self.assertIn("单位水平力地面处位移", xml)
        self.assertIn("标准组合水平力下地面处位移", xml)
        self.assertNotIn("因换算埋深αh小于2.4", xml)

    def test_definition_scope_exceedance_is_not_written_to_report(self) -> None:
        base = sample_grouted()
        common = replace(base.common, diameter_m=0.35, embedment_m=5.1)
        soils = (replace(base.soils[1], thickness_m=5.1),)
        result = calculate(replace(base, common=common, soils=soils))
        xml, _styles, _all_xml, _media, *_checks = self._generate_and_read_xml(result)
        self.assertNotIn("超出微型短桩", xml)
        self.assertNotIn("定义范围", xml)


if __name__ == "__main__":
    unittest.main()
